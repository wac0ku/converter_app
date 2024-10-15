# Author: Leon Gajtner
# Datum: 15.10.2024
# PDF Magic Converter

from header import PDFConverterInterface, DropAreaInterface
from PyQt5.QtWidgets import QTextEdit, QPushButton, QProgressBar, QFileDialog, QVBoxLayout, QLabel
from PyQt5.QtGui import QDragEnterEvent, QDropEvent
from PyQt5.QtCore import Qt 
from pdf2docx import Converter
import docx
import os
import logging

class PDFConverter(PDFConverterInterface):
    def run(self):
        total_files = len(self.pdf_files)
        for index, pdf_file in enumerate(self.pdf_files, start=1):
            try:
                base_name = os.path.basename(pdf_file)
                docx_file = os.path.join(self.output_directory, os.path.splitext(base_name)[0] + '.docx')
                cv = Converter(pdf_file)
                cv.convert(docx_file)
                cv.close()

                # Formatierung anpassen
                doc = docx.Document(docx_file)
                for paragraph in doc.paragraphs:
                    paragraph.style.font.name = 'Arial'
                    paragraph.style.font.size = docx.shared.Pt(11)

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph.style.font.name = 'Arial'
                                paragraph.style.font.size = docx.shared.Pt(11)

                doc.save(docx_file)
                success_message = f"Erfolgreich konvertiert: {pdf_file}"
                self.update_log.emit(success_message)
                logging.info(success_message)
            except Exception as e:
                error_message = f"Fehler bei der Konvertierung von {pdf_file}: {str(e)}"
                self.update_log.emit(error_message)
                logging.error(error_message)
            
            progress = int((index / total_files) * 100)
            self.update_progress.emit(progress)

class DropArea(DropAreaInterface):
    def __init__(self):
        super().__init__()
        layout = QVBoxLayout()
        self.label = QLabel("PDF-Dateien hier hineinziehen")
        self.label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.label)
        self.setLayout(layout)

    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            event.accept()
        else:
            event.ignore()

    def dragMoveEvent(self, event):
        if event.mimeData().hasUrls():
            event.setDropAction(Qt.CopyAction)
            event.accept()
        else:
            event.ignore()

    def dropEvent(self, event: QDropEvent):
        files = [url.toLocalFile() for url in event.mimeData().urls() if url.toLocalFile().lower().endswith('.pdf')]
        if files:
            event.setDropAction(Qt.CopyAction)
            event.accept()
            self.files_dropped.emit(files)
        else:
            event.ignore()